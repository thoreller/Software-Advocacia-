import sqlite3
from docx import Document
from docx.shared import Pt 
import os
import sys 
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import requests
import locale

# --- CONFIGURAÇÃO INICIAL ---
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')

# --- FUNÇÃO CORRETA PARA ENCONTRAR ARQUIVOS NO EXECUTÁVEL ---
def resource_path(relative_path):
    """ Obtém o caminho absoluto para o recurso, funciona para dev e para PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- CONFIGURAÇÃO DE CAMINHOS ---
MODELO_PROCURACAO = resource_path("modelo_procuracao.docx")
MODELO_HIPOSSUFICIENCIA = resource_path("modelo_hipossuficiencia.docx")
MODELO_HONORARIOS = resource_path("modelo_honorarios.docx")
DB_NAME = os.path.join(os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__)), "cadastros_clientes.db")

# --- FUNÇÕES DE FORMATAÇÃO ---
def formatar_cpf(event=None):
    texto_atual = entry_cpf.get()
    numeros = ''.join(filter(str.isdigit, texto_atual))[:11]
    texto_formatado = ""
    if len(numeros) > 9:
        texto_formatado = f"{numeros[:3]}.{numeros[3:6]}.{numeros[6:9]}-{numeros[9:]}"
    elif len(numeros) > 6:
        texto_formatado = f"{numeros[:3]}.{numeros[3:6]}.{numeros[6:]}"
    elif len(numeros) > 3:
        texto_formatado = f"{numeros[:3]}.{numeros[3:]}"
    else:
        texto_formatado = numeros
    if texto_atual != texto_formatado:
        entry_cpf.delete(0, tk.END)
        entry_cpf.insert(0, texto_formatado)
        entry_cpf.icursor(len(texto_formatado))

def formatar_data(event=None):
    texto_atual = entry_nascimento.get()
    numeros = ''.join(filter(str.isdigit, texto_atual))[:8]
    texto_formatado = ""
    if len(numeros) > 4:
        texto_formatado = f"{numeros[:2]}/{numeros[2:4]}/{numeros[4:]}"
    elif len(numeros) > 2:
        texto_formatado = f"{numeros[:2]}/{numeros[2:]}"
    else:
        texto_formatado = numeros
    if texto_atual != texto_formatado:
        entry_nascimento.delete(0, tk.END)
        entry_nascimento.insert(0, texto_formatado)
        entry_nascimento.icursor(len(texto_formatado))

# --- FUNÇÕES PRINCIPAIS DO PROGRAMA ---
def verificar_e_atualizar_banco():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS clientes (id INTEGER PRIMARY KEY AUTOINCREMENT, nome_completo TEXT NOT NULL, cpf TEXT NOT NULL)''')
    colunas_necessarias = [
        ('genero', 'TEXT'), ('estado_civil', 'TEXT'), ('profissao', 'TEXT'), ('data_nascimento', 'TEXT'),
        ('nome_mae', 'TEXT'), ('nome_pai', 'TEXT'), ('cep', 'TEXT'), ('logradouro', 'TEXT'),
        ('numero', 'TEXT'), ('complemento', 'TEXT'), ('bairro', 'TEXT'), ('cidade', 'TEXT'),
        ('uf', 'TEXT'), ('cidade_assinatura', 'TEXT'), ('data_documento', 'TEXT')
    ]
    c.execute("PRAGMA table_info(clientes)")
    colunas_existentes = [info[1] for info in c.fetchall()]
    for coluna, tipo in colunas_necessarias:
        if coluna not in colunas_existentes:
            c.execute(f"ALTER TABLE clientes ADD COLUMN {coluna} {tipo}")
    conn.commit()
    conn.close()

def buscar_cep():
    cep = entry_cep.get().replace("-", "").replace(".", "").strip()
    if len(cep) != 8:
        messagebox.showerror("Erro de CEP", "O CEP deve conter 8 dígitos.")
        return
    try:
        response = requests.get(f"https://viacep.com.br/ws/{cep}/json/")
        response.raise_for_status()
        dados_cep = response.json()
        if dados_cep.get("erro"):
            messagebox.showerror("Erro de CEP", "CEP não encontrado.")
            return
        for entry in [entry_logradouro, entry_bairro, entry_cidade, entry_uf]:
            entry.delete(0, tk.END)
        entry_logradouro.insert(0, dados_cep.get('logradouro', ''))
        entry_bairro.insert(0, dados_cep.get('bairro', ''))
        entry_cidade.insert(0, dados_cep.get('localidade', ''))
        entry_uf.insert(0, dados_cep.get('uf', ''))
        entry_numero.focus_set()
    except requests.exceptions.RequestException:
        messagebox.showerror("Erro de Rede", "Não foi possível conectar à API de CEP.")
    except Exception as e:
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro ao buscar o CEP: {e}")

def gerar_documentos(dados_cliente):
    pasta_selecionada = filedialog.askdirectory(title="Selecione a pasta para salvar os documentos")
    if not pasta_selecionada:
        messagebox.showinfo("Operação Cancelada", "A geração de documentos foi cancelada.")
        return

    # --- LÓGICA DE GÊNERO E TEXTOS (COMPLETA E CORRIGIDA) ---
    genero = dados_cliente['genero']
    profissao_original = dados_cliente['profissao']
    estado_civil_original = dados_cliente['estado_civil']
    
    base_estado_civil = estado_civil_original.replace('(a)', '')
    if genero == "Feminino":
        nacionalidade = "brasileira"
        filho_a = "filha"
        nascido_a = "nascida"
        inscrito_a = "inscrita"
        denominado_a = "denominada"
        estado_civil_ajustado = base_estado_civil[:-1] + 'a' if base_estado_civil.endswith('o') else base_estado_civil
        profissao_ajustada = profissao_original.replace("PROFESSOR DE", "PROFESSORA DE")
        profissao_ajustada = profissao_ajustada.replace("ASSISTENTE TÉCNICO", "ASSISTENTE TÉCNICA")
        profissao_ajustada = profissao_ajustada.replace("FUNCIONÁRIO PÚBLICO", "FUNCIONÁRIA PÚBLICA")
    else: # Masculino
        nacionalidade = "brasileiro"
        filho_a = "filho"
        nascido_a = "nascido"
        inscrito_a = "inscrito"
        denominado_a = "denominado"
        estado_civil_ajustado = base_estado_civil if not base_estado_civil.endswith('o') else base_estado_civil
        profissao_ajustada = profissao_original

    # --- MONTAGEM DOS TEXTOS ---
    endereco_completo = (
        f"{dados_cliente['logradouro']}, Nº {dados_cliente['numero']}"
        f"{', ' + dados_cliente['complemento'] if dados_cliente['complemento'] else ''}, "
        f"{dados_cliente['bairro']}, {dados_cliente['cidade']}/{dados_cliente['uf']}, CEP: {dados_cliente['cep']}"
    )
    
    local_data = f"{dados_cliente['cidade_assinatura']}, {dados_cliente['data_documento'].lower()}."

    documentos_a_gerar = [(MODELO_PROCURACAO, "Procuracao"), (MODELO_HIPOSSUFICIENCIA, "Hipossuficiencia"), (MODELO_HONORARIOS, "Honorarios")]
    primeiro_nome = dados_cliente['nome_completo'].split(" ")[0]
    cpf_numeros = ''.join(filter(str.isdigit, dados_cliente['cpf']))
    nome_base_arquivo = os.path.join(pasta_selecionada, f"{primeiro_nome}_{cpf_numeros}")
    
    try:
        for modelo_path, sufixo in documentos_a_gerar:
            if not os.path.exists(modelo_path):
                messagebox.showerror("Erro de Arquivo", f"Arquivo modelo não encontrado:\n{modelo_path}")
                return
                
            doc = Document(modelo_path)
            
            for p in doc.paragraphs:
                # Lógica de substituição com formatação
                if '{{' in p.text:
                    # PROCURAÇÃO E HIPOSSUFICIÊNCIA (Campos em negrito)
                    if '{{NOME_COMPLETO}}' in p.text:
                        p.clear()
                        p.add_run('NOME: ').bold = False
                        p.add_run(dados_cliente['nome_completo'].upper()).bold = True
                    elif '{{CPF}}' in p.text:
                        p.clear()
                        p.add_run('CPF: ').bold = False
                        p.add_run(dados_cliente['cpf']).bold = True
                    elif '{{ESTADO_CIVIL}}' in p.text:
                        p.clear()
                        p.add_run('ESTADO CIVIL: ').bold = False
                        p.add_run(estado_civil_ajustado.upper()).bold = True
                    elif '{{PROFISSAO}}' in p.text:
                        p.clear()
                        p.add_run('PROFISSÃO: ').bold = False
                        p.add_run(profissao_ajustada.upper()).bold = True
                    elif '{{ENDERECO_COMPLETO}}' in p.text:
                        p.clear()
                        p.add_run('ENDEREÇO RESIDENCIAL: ').bold = False
                        p.add_run(endereco_completo.upper()).bold = True
                    
                    # HONORÁRIOS (Qualificação complexa)
                    elif '{{QUALIFICACAO_COMPLETA}}' in p.text:
                        p.clear()
                        p.add_run('– ')
                        run_nome = p.add_run(dados_cliente['nome_completo'].upper())
                        run_nome.bold = True
                        run_nome.font.underline = True
                        
                        resto_qualificacao = (
                            f", {nacionalidade}, {estado_civil_ajustado.lower()}, "
                            f"{profissao_ajustada.lower()}, {nascido_a} em {dados_cliente['data_nascimento']}, "
                            f"{inscrito_a} no CPF sob o n.º {dados_cliente['cpf']}, {filho_a} de {dados_cliente['nome_mae']} e {dados_cliente['nome_pai']}, "
                            f"residente e domiciliado(a) na {endereco_completo}"
                        )
                        p.add_run(resto_qualificacao)
                        
                        p.add_run(f", que por força do presente contrato passa a ser {denominado_a}(a) ")
                        run_contratante = p.add_run('Contratante')
                        run_contratante.bold = True
                        run_contratante.font.underline = True
                        p.add_run(';')

                    # DATA (para todos os documentos)
                    elif '{{LOCAL_DATA}}' in p.text:
                        p.text = p.text.replace('{{LOCAL_DATA}}', local_data)

            nome_arquivo_saida = f"{nome_base_arquivo}_{sufixo}.docx"
            doc.save(nome_arquivo_saida)
        
        messagebox.showinfo("Sucesso", f"Documentos gerados com sucesso na pasta:\n{pasta_selecionada}")

    except Exception as e:
        messagebox.showerror("Erro ao Gerar Documentos", f"Ocorreu um erro: {e}")

def cadastrar_e_gerar():
    dados_cliente = {
        "nome_completo": entry_nome.get(), "cpf": entry_cpf.get(),
        "genero": combo_genero.get(), "estado_civil": combo_estado_civil.get(),
        "profissao": combo_profissao.get(), "data_nascimento": entry_nascimento.get(),
        "nome_mae": entry_mae.get(), "nome_pai": entry_pai.get(), "cep": entry_cep.get(),
        "logradouro": entry_logradouro.get(), "numero": entry_numero.get(),
        "complemento": entry_complemento.get(), "bairro": entry_bairro.get(),
        "cidade": entry_cidade.get(), "uf": entry_uf.get(),
        "cidade_assinatura": entry_cidade_assinatura.get(), "data_documento": entry_data.get()
    }
    if not all([dados_cliente['nome_completo'], dados_cliente['cpf'], dados_cliente['genero']]):
        messagebox.showwarning("Atenção", "Os campos 'Nome Completo', 'CPF' e 'Gênero' são obrigatórios.")
        return
        
    gerar_documentos(dados_cliente)
    
    try:
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("""
            INSERT INTO clientes (
                nome_completo, cpf, genero, estado_civil, profissao, data_nascimento, 
                nome_mae, nome_pai, cep, logradouro, numero, complemento, 
                bairro, cidade, uf, cidade_assinatura, data_documento
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, tuple(dados_cliente.values()))
        conn.commit()
        conn.close()
    except Exception as e:
        messagebox.showerror("Erro de Banco de Dados", f"Não foi possível salvar o cliente: {e}")
        return
    
    # Limpa os campos e atualiza a lista de consulta
    widgets_a_limpar = [
        entry_nome, entry_cpf, entry_nascimento, entry_mae, entry_pai,
        entry_cep, entry_logradouro, entry_numero, entry_complemento,
        entry_bairro, entry_cidade, entry_uf, entry_cidade_assinatura, entry_data
    ]
    combos_a_limpar = [combo_genero, combo_estado_civil, combo_profissao]

    for widget in widgets_a_limpar:
        widget.delete(0, tk.END)

    for combo in combos_a_limpar:
        combo.set('')
    
    entry_nome.focus_set()
    entry_data.insert(0, datetime.now().strftime("%d de %B de %Y"))
    entry_cidade_assinatura.insert(0, "Governador Valadares/MG")
    carregar_clientes() # Atualiza a aba de consulta

def carregar_clientes():
    """Carrega os clientes do banco de dados e exibe na aba de consulta."""
    for i in tree_consulta.get_children():
        tree_consulta.delete(i)
    
    try:
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute("SELECT id, nome_completo, cpf, data_nascimento FROM clientes ORDER BY nome_completo")
        for row in c.fetchall():
            tree_consulta.insert("", "end", values=row)
        conn.close()
    except Exception as e:
        messagebox.showerror("Erro de Banco de Dados", f"Não foi possível carregar os clientes: {e}")

# --- INTERFACE GRÁFICA (Tkinter) ---
verificar_e_atualizar_banco()
janela = tk.Tk()
janela.title("Gerador de Documentos para Advocacia v10.0 (Final)")
janela.geometry("650x700")

# --- Criação das Abas ---
notebook = ttk.Notebook(janela)
notebook.pack(pady=10, padx=10, fill="both", expand=True)

tab_cadastro = ttk.Frame(notebook)
tab_consulta = ttk.Frame(notebook)

notebook.add(tab_cadastro, text='Cadastrar Cliente')
notebook.add(tab_consulta, text='Consultar Cadastros')

# --- ABA 1: CADASTRO DE CLIENTE ---
frame_cadastro = tk.Frame(tab_cadastro, padx=10, pady=10)
frame_cadastro.pack(fill="both", expand=True)

frame_dados_pessoais = tk.LabelFrame(frame_cadastro, text="Dados Pessoais", padx=10, pady=10)
frame_dados_pessoais.pack(fill="x", expand=True, pady=5)

tk.Label(frame_dados_pessoais, text="Nome Completo:").grid(row=0, column=0, sticky="w", pady=2)
entry_nome = tk.Entry(frame_dados_pessoais, width=50)
entry_nome.grid(row=0, column=1, columnspan=3, sticky="we")
tk.Label(frame_dados_pessoais, text="CPF:").grid(row=1, column=0, sticky="w", pady=2)
entry_cpf = tk.Entry(frame_dados_pessoais)
entry_cpf.grid(row=1, column=1, sticky="we")
entry_cpf.bind("<KeyRelease>", formatar_cpf)
tk.Label(frame_dados_pessoais, text="Data de Nasc.:").grid(row=1, column=2, sticky="w", pady=2, padx=(10,0))
entry_nascimento = tk.Entry(frame_dados_pessoais)
entry_nascimento.grid(row=1, column=3, sticky="we")
entry_nascimento.bind("<KeyRelease>", formatar_data)
tk.Label(frame_dados_pessoais, text="Gênero:").grid(row=2, column=0, sticky="w", pady=2)
combo_genero = ttk.Combobox(frame_dados_pessoais, values=['Masculino', 'Feminino'], state="readonly")
combo_genero.grid(row=2, column=1, sticky="we")
tk.Label(frame_dados_pessoais, text="Estado Civil:").grid(row=2, column=2, sticky="w", pady=2, padx=(10,0))
opcoes_estado_civil = ['Solteiro(a)', 'Casado(a)', 'Divorciado(a)', 'Viúvo(a)', 'União Estável']
combo_estado_civil = ttk.Combobox(frame_dados_pessoais, values=opcoes_estado_civil)
combo_estado_civil.grid(row=2, column=3, sticky="we")
tk.Label(frame_dados_pessoais, text="Profissão:").grid(row=3, column=0, sticky="w", pady=2)
opcoes_profissao = ['PROFESSOR DE EDUCAÇÃO BÁSICA', 'ASSISTENTE TÉCNICO DE EDUCAÇÃO BÁSICA', 'AUXILIAR DE SERVIÇOS DE EDUCAÇÃO BÁSICA', 'FUNCIONÁRIO PÚBLICO']
combo_profissao = ttk.Combobox(frame_dados_pessoais, values=opcoes_profissao, state="readonly")
combo_profissao.grid(row=3, column=1, columnspan=3, sticky="we")
tk.Label(frame_dados_pessoais, text="Nome da Mãe:").grid(row=4, column=0, sticky="w", pady=2)
entry_mae = tk.Entry(frame_dados_pessoais)
entry_mae.grid(row=4, column=1, columnspan=3, sticky="we")
tk.Label(frame_dados_pessoais, text="Nome do Pai:").grid(row=5, column=0, sticky="w", pady=2)
entry_pai = tk.Entry(frame_dados_pessoais)
entry_pai.grid(row=5, column=1, columnspan=3, sticky="we")

frame_endereco = tk.LabelFrame(frame_cadastro, text="Endereço", padx=10, pady=10)
frame_endereco.pack(fill="x", expand=True, pady=5)
# ... Widgets de endereço ...
tk.Label(frame_endereco, text="CEP:").grid(row=0, column=0, sticky="w", pady=2)
entry_cep = tk.Entry(frame_endereco)
entry_cep.grid(row=0, column=1, sticky="we")
btn_buscar_cep = tk.Button(frame_endereco, text="Buscar", command=buscar_cep)
btn_buscar_cep.grid(row=0, column=2, padx=(5,0))
tk.Label(frame_endereco, text="Logradouro (Rua/Av):").grid(row=1, column=0, sticky="w", pady=2)
entry_logradouro = tk.Entry(frame_endereco, width=50)
entry_logradouro.grid(row=1, column=1, columnspan=3, sticky="we")
tk.Label(frame_endereco, text="Número:").grid(row=2, column=0, sticky="w", pady=2)
entry_numero = tk.Entry(frame_endereco)
entry_numero.grid(row=2, column=1, sticky="we")
tk.Label(frame_endereco, text="Complemento:").grid(row=2, column=2, sticky="w", pady=2, padx=(10,0))
entry_complemento = tk.Entry(frame_endereco)
entry_complemento.grid(row=2, column=3, sticky="we")
tk.Label(frame_endereco, text="Bairro:").grid(row=3, column=0, sticky="w", pady=2)
entry_bairro = tk.Entry(frame_endereco)
entry_bairro.grid(row=3, column=1, columnspan=3, sticky="we")
tk.Label(frame_endereco, text="Cidade:").grid(row=4, column=0, sticky="w", pady=2)
entry_cidade = tk.Entry(frame_endereco)
entry_cidade.grid(row=4, column=1, sticky="we")
tk.Label(frame_endereco, text="UF:").grid(row=4, column=2, sticky="w", pady=2, padx=(10,0))
entry_uf = tk.Entry(frame_endereco)
entry_uf.grid(row=4, column=3, sticky="we")

frame_assinatura = tk.LabelFrame(frame_cadastro, text="Dados do Documento", padx=10, pady=10)
frame_assinatura.pack(fill="x", expand=True, pady=5)
# ... Widgets de assinatura ...
tk.Label(frame_assinatura, text="Cidade da Assinatura:").grid(row=0, column=0, sticky="w", pady=2)
entry_cidade_assinatura = tk.Entry(frame_assinatura)
entry_cidade_assinatura.grid(row=0, column=1, sticky="we")
entry_cidade_assinatura.insert(0, "Governador Valadares/MG")
tk.Label(frame_assinatura, text="Data do Documento:").grid(row=0, column=2, sticky="w", pady=2, padx=(10,0))
entry_data = tk.Entry(frame_assinatura)
entry_data.grid(row=0, column=3, sticky="we")
entry_data.insert(0, datetime.now().strftime("%d de %B de %Y"))

btn_cadastrar = tk.Button(frame_cadastro, text="Cadastrar Cliente e Gerar Documentos", command=cadastrar_e_gerar, bg="#28a745", fg="white", font=("Helvetica", 12, "bold"), pady=10)
btn_cadastrar.pack(pady=20)

# --- ABA 2: CONSULTA DE CLIENTES ---
frame_consulta = tk.Frame(tab_consulta, padx=10, pady=10)
frame_consulta.pack(fill="both", expand=True)

frame_botoes_consulta = tk.Frame(frame_consulta)
frame_botoes_consulta.pack(fill="x", pady=5)
btn_atualizar_lista = tk.Button(frame_botoes_consulta, text="Atualizar Lista", command=carregar_clientes)
btn_atualizar_lista.pack(side="left")

# Criação da tabela (Treeview)
cols = ('ID', 'Nome Completo', 'CPF', 'Data de Nascimento')
tree_consulta = ttk.Treeview(frame_consulta, columns=cols, show='headings')
for col in cols:
    tree_consulta.heading(col, text=col)
tree_consulta.column('ID', width=50)
tree_consulta.column('Nome Completo', width=250)
tree_consulta.column('CPF', width=120)
tree_consulta.column('Data de Nascimento', width=120)

tree_consulta.pack(fill="both", expand=True)

# Carrega os clientes na inicialização do programa
carregar_clientes()
janela.mainloop()